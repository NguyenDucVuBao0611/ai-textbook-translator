import os
import hashlib
import json
import asyncio
from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from docx import Document as DocxDocument

from backend.database import SessionLocal, init_db, Document as DBDocument, AINote as DBAINote
from backend.file_processor import translate_docx, extract_pdf_text_by_pages, translate_txt
from backend.chunking import chunk_text
from backend.translate import TranslationCoordinator

# Ensure folders exist
UPLOAD_DIR = "uploads"
TRANSLATED_DIR = "translated"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TRANSLATED_DIR, exist_ok=True)

app = FastAPI(title="AI Textbook Translator API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global dict to store streaming progress
# Format: {document_id: {"status": "processing/completed/failed", "pages": {page_num: text}}}
translation_progress = {}

# DB Dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Translation Coordinator instance (loaded once on startup)
translator = None

@app.on_event("startup")
def startup_event():
    global translator
    init_db()  # Ensure SQLite tables exist
    translator = TranslationCoordinator()

def run_translation_in_background(document_id: str, file_path: str, file_type: str, mode: str):
    global translator, translation_progress
    translation_progress[document_id] = {"status": "processing", "pages": {}}
    
    output_filename = f"translated_{document_id}.docx"
    output_path = os.path.join(TRANSLATED_DIR, output_filename)
    
    try:
        def translate_paragraph(text: str) -> str:
            # Helper to chunk and translate text
            chunks = chunk_text(text, max_words=120)
            translated_chunks = []
            for chunk in chunks:
                translated_chunks.append(translator.translate_sentence(chunk, mode=mode))
            return " ".join(translated_chunks)
            
        if file_type == "docx":
            # Word file translation
            translate_docx(file_path, translate_paragraph, output_path)
            translation_progress[document_id]["pages"][1] = "Dịch tài liệu Word thành công."
            translation_progress[document_id]["status"] = "completed"
            
        elif file_type == "txt":
            # Plain text file translation
            translate_txt(file_path, translate_paragraph, output_path)
            translation_progress[document_id]["pages"][1] = "Dịch tài liệu văn bản thành công."
            translation_progress[document_id]["status"] = "completed"
            
        elif file_type == "pdf":
            # PDF file translation page-by-page
            pages = extract_pdf_text_by_pages(file_path)
            translated_pages = []
            
            for idx, page_text in enumerate(pages, 1):
                if page_text.strip():
                    translated_page = translate_paragraph(page_text)
                else:
                    translated_page = ""
                
                # Store in progress for streaming
                translation_progress[document_id]["pages"][idx] = translated_page
                translated_pages.append(translated_page)
                
            # For PDF, save the translation output as a DOCX file for easy download
            doc = DocxDocument()
            for idx, page_content in enumerate(translated_pages, 1):
                doc.add_heading(f"Trang {idx}", level=2)
                doc.add_paragraph(page_content)
            doc.save(output_path)
            
            translation_progress[document_id]["status"] = "completed"
            
    except Exception as e:
        print(f"Error in background translation for {document_id}: {e}")
        translation_progress[document_id]["status"] = "failed"

@app.post("/api/upload")
async def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    mode: str = Form("academic"),
    db: Session = Depends(get_db)
):
    # Determine file extension
    filename = file.filename
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    if ext not in ["docx", "pdf", "txt"]:
        raise HTTPException(status_code=400, detail="Only .docx, .pdf, and .txt files are supported.")
        
    # Read file content to generate MD5 hash
    contents = await file.read()
    document_id = hashlib.md5(contents).hexdigest()
    
    # Save the file to disk
    file_path = os.path.join(UPLOAD_DIR, f"{document_id}.{ext}")
    with open(file_path, "wb") as f:
        f.write(contents)
        
    # Verify if document is already in SQLite
    db_doc = db.query(DBDocument).filter(DBDocument.id == document_id).first()
    if not db_doc:
        db_doc = DBDocument(id=document_id, filename=filename)
        db.add(db_doc)
        db.commit()
        
    # Start background translation task
    background_tasks.add_task(run_translation_in_background, document_id, file_path, ext, mode)
    
    return {
        "document_id": document_id,
        "filename": filename,
        "status": "processing",
        "message": "File uploaded successfully. Translation started."
    }

@app.get("/api/stream/{document_id}")
async def stream_progress(document_id: str):
    """
    Server-Sent Events (SSE) endpoint to stream translation progress of the document page-by-page.
    """
    async def event_generator():
        sent_pages = set()
        
        while True:
            # Check if document exists in global progress dictionary
            if document_id not in translation_progress:
                yield f"data: {json.dumps({'error': 'Document not found'})}\n\n"
                break
                
            doc_data = translation_progress[document_id]
            status = doc_data["status"]
            pages = doc_data["pages"]
            
            # Send newly translated pages
            for page_num in sorted(pages.keys()):
                if page_num not in sent_pages:
                    payload = {
                        "page": page_num,
                        "text": pages[page_num],
                        "status": status
                    }
                    yield f"data: {json.dumps(payload)}\n\n"
                    sent_pages.add(page_num)
                    
            if status == "completed":
                yield f"data: {json.dumps({'status': 'completed', 'message': 'Translation finished'})}\n\n"
                break
            elif status == "failed":
                yield f"data: {json.dumps({'status': 'failed', 'message': 'Translation failed'})}\n\n"
                break
                
            await asyncio.sleep(1.0)
            
    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.get("/api/download/{document_id}")
async def download_translated_file(document_id: str):
    output_path = os.path.join(TRANSLATED_DIR, f"translated_{document_id}.docx")
    if not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="Translated file not found.")
    return FileResponse(output_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"translated_{document_id}.docx")
