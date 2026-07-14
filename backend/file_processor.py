import os
from docx import Document

def translate_docx(file_path: str, translate_func, output_path: str):
    """
    Reads a Word file, translates its paragraphs and table cells using the callback translate_func,
    and writes the translated content to output_path while preserving paragraph/cell structure.
    """
    doc = Document(file_path)
    
    # 1. Translate all paragraphs
    print(f"Translating DOCX paragraphs for: {os.path.basename(file_path)}")
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Translate paragraph text
            translated_text = translate_func(paragraph.text)
            paragraph.text = translated_text
            
    # 2. Translate all tables cell-by-cell
    print(f"Translating DOCX tables for: {os.path.basename(file_path)}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Avoid translating repeated merged cells multiple times
                if cell.text.strip():
                    translated_cell = translate_func(cell.text)
                    cell.text = translated_cell
                    
    doc.save(output_path)
    print(f"Saved translated DOCX to: {output_path}")

def extract_pdf_text_by_pages(file_path: str) -> list:
    """
    Extracts text page-by-page from a PDF file using pdfplumber.
    Returns a list of strings, where each string represents a page's text.
    """
    import pdfplumber
    pages_text = []
    print(f"Extracting PDF text from: {os.path.basename(file_path)}")
    with pdfplumber.open(file_path) as pdf:
        for idx, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            pages_text.append(text)
    return pages_text

def translate_txt(file_path: str, translate_func, output_path: str):
    """
    Translates a plain text file line by line and writes to output_path.
    """
    print(f"Translating TXT file: {os.path.basename(file_path)}")
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    translated_lines = []
    for line in lines:
        if line.strip():
            translated_lines.append(translate_func(line) + "\n")
        else:
            translated_lines.append("\n")
            
    with open(output_path, "w", encoding="utf-8") as f:
        f.writelines(translated_lines)
    print(f"Saved translated TXT to: {output_path}")
